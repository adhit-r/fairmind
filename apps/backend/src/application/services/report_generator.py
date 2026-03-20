"""
Report Generator Service
Generates PDF and DOCX reports for bias evaluations, compliance assessments, and governance audits.
"""

from io import BytesIO
from datetime import datetime
from typing import Optional, List, Dict, Any

from reportlab.lib.pagesizes import letter, A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch, mm
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak, Table, TableStyle, Image
from reportlab.lib import colors
from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_RIGHT, TA_JUSTIFY

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ─── Data Models ────────────────────────────────────────────────────────────

class ComplianceStatus:
    def __init__(
        self,
        gdpr_compliant: bool = False,
        ai_act_compliant: bool = False,
        fairness_score: Optional[float] = None,
    ):
        self.gdpr_compliant = gdpr_compliant
        self.ai_act_compliant = ai_act_compliant
        self.fairness_score = fairness_score


class EvaluationSummary:
    def __init__(
        self,
        total_tests: int,
        tests_passed: int,
        tests_failed: int,
        overall_bias_rate: float,
        evaluation_time: Optional[str] = None,
    ):
        self.total_tests = total_tests
        self.tests_passed = tests_passed
        self.tests_failed = tests_failed
        self.overall_bias_rate = overall_bias_rate
        self.evaluation_time = evaluation_time


class BiasReportData:
    def __init__(
        self,
        timestamp: str,
        model_type: str,
        model_description: Optional[str] = None,
        evaluation_summary: Optional[EvaluationSummary] = None,
        overall_risk: str = "medium",
        risk_factors: Optional[List[str]] = None,
        recommendations: Optional[List[str]] = None,
        compliance_status: Optional[ComplianceStatus] = None,
        explainability_insights: Optional[List[str]] = None,
    ):
        self.timestamp = timestamp
        self.model_type = model_type
        self.model_description = model_description
        self.evaluation_summary = evaluation_summary
        self.overall_risk = overall_risk
        self.risk_factors = risk_factors or []
        self.recommendations = recommendations or []
        self.compliance_status = compliance_status
        self.explainability_insights = explainability_insights or []


# ─── ReportLab PDF Generator ────────────────────────────────────────────────

class PDFReportGenerator:
    """Generate professional PDF reports using ReportLab."""

    COLORS = {
        "low": colors.HexColor("#22c55e"),
        "medium": colors.HexColor("#eab308"),
        "high": colors.HexColor("#f97316"),
        "critical": colors.HexColor("#ef4444"),
        "primary": colors.HexColor("#1f2937"),
        "accent": colors.HexColor("#ff6b35"),
        "text": colors.HexColor("#111827"),
        "muted": colors.HexColor("#6b7280"),
    }

    def generate(self, data: BiasReportData) -> bytes:
        """Generate PDF report and return as bytes."""
        buffer = BytesIO()
        doc = SimpleDocTemplate(
            buffer,
            pagesize=letter,
            rightMargin=0.75 * inch,
            leftMargin=0.75 * inch,
            topMargin=0.75 * inch,
            bottomMargin=0.75 * inch,
        )

        # Build content
        story = []
        styles = self._get_styles()

        # Header
        story.extend(self._build_header(styles, data))

        # Executive Summary
        story.append(Spacer(1, 0.3 * inch))
        story.extend(self._build_executive_summary(styles, data))

        # Evaluation Summary
        story.append(PageBreak())
        story.extend(self._build_evaluation_summary(styles, data))

        # Compliance Status
        story.append(Spacer(1, 0.3 * inch))
        story.extend(self._build_compliance_status(styles, data))

        # Risk Factors
        if data.risk_factors:
            story.append(Spacer(1, 0.3 * inch))
            story.extend(self._build_risk_factors(styles, data))

        # Recommendations
        story.append(Spacer(1, 0.3 * inch))
        story.extend(self._build_recommendations(styles, data))

        # Footer
        story.append(Spacer(1, 0.5 * inch))
        story.extend(self._build_footer(styles, data))

        # Build PDF
        doc.build(story)
        buffer.seek(0)
        return buffer.getvalue()

    def _get_styles(self):
        """Get or create document styles."""
        styles = getSampleStyleSheet()

        # Title
        styles.add(
            ParagraphStyle(
                name="CustomTitle",
                parent=styles["Heading1"],
                fontSize=28,
                textColor=self.COLORS["primary"],
                spaceAfter=6,
                fontName="Helvetica-Bold",
            )
        )

        # Heading
        styles.add(
            ParagraphStyle(
                name="CustomHeading",
                parent=styles["Heading2"],
                fontSize=14,
                textColor=self.COLORS["primary"],
                spaceAfter=12,
                fontName="Helvetica-Bold",
                borderPadding=12,
                backColor=colors.HexColor("#f3f4f6"),
            )
        )

        # Body
        styles.add(
            ParagraphStyle(
                name="CustomBody",
                parent=styles["BodyText"],
                fontSize=10,
                textColor=self.COLORS["text"],
                spaceAfter=12,
            )
        )

        return styles

    def _build_header(self, styles, data):
        """Build report header with FairMind branding."""
        elements = []
        elements.append(Paragraph("FairMind", styles["CustomTitle"]))
        elements.append(
            Paragraph(
                "AI Governance & Bias Detection Platform",
                styles["Normal"],
            )
        )
        elements.append(Paragraph("Bias Evaluation Report", styles["Heading2"]))
        elements.append(
            Paragraph(
                f"Generated: {datetime.fromisoformat(data.timestamp).strftime('%B %d, %Y at %H:%M UTC')}",
                styles["Normal"],
            )
        )
        return elements

    def _build_executive_summary(self, styles, data):
        """Build executive summary section."""
        elements = []
        elements.append(Paragraph("Executive Summary", styles["CustomHeading"]))

        # Risk level
        risk_color = self.COLORS.get(data.overall_risk, self.COLORS["medium"])
        elements.append(
            Paragraph(
                f"<b>Overall Risk Level:</b> <font color='{risk_color.hexval()}'>{data.overall_risk.upper()}</font>",
                styles["Normal"],
            )
        )

        # Key metrics
        if data.evaluation_summary:
            bias_rate = (data.evaluation_summary.overall_bias_rate * 100)
            pass_rate = (data.evaluation_summary.tests_passed / max(1, data.evaluation_summary.total_tests)) * 100
            elements.append(
                Paragraph(
                    f"<b>Bias Rate:</b> {bias_rate:.1f}% | <b>Tests Passed:</b> {data.evaluation_summary.tests_passed}/{data.evaluation_summary.total_tests} ({pass_rate:.0f}%)",
                    styles["Normal"],
                )
            )

        return elements

    def _build_evaluation_summary(self, styles, data):
        """Build evaluation summary table."""
        elements = []
        elements.append(Paragraph("Evaluation Summary", styles["CustomHeading"]))

        if data.evaluation_summary:
            table_data = [
                ["Metric", "Value"],
                ["Total Tests", str(data.evaluation_summary.total_tests)],
                ["Tests Passed", str(data.evaluation_summary.tests_passed)],
                ["Tests Failed", str(data.evaluation_summary.tests_failed)],
                ["Overall Bias Rate", f"{data.evaluation_summary.overall_bias_rate * 100:.1f}%"],
            ]

            table = Table(table_data, colWidths=[2.5 * inch, 2.5 * inch])
            table.setStyle(
                TableStyle(
                    [
                        ("BACKGROUND", (0, 0), (-1, 0), self.COLORS["primary"]),
                        ("TEXTCOLOR", (0, 0), (-1, 0), colors.whitesmoke),
                        ("ALIGN", (0, 0), (-1, -1), "CENTER"),
                        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                        ("FONTSIZE", (0, 0), (-1, 0), 12),
                        ("BOTTOMPADDING", (0, 0), (-1, 0), 12),
                        ("GRID", (0, 0), (-1, -1), 1, colors.grey),
                        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#f9fafb")]),
                    ]
                )
            )
            elements.append(table)

        return elements

    def _build_compliance_status(self, styles, data):
        """Build compliance status section."""
        elements = []
        elements.append(Paragraph("Compliance Status", styles["CustomHeading"]))

        if data.compliance_status:
            gdpr_status = "✓ COMPLIANT" if data.compliance_status.gdpr_compliant else "✗ NON-COMPLIANT"
            ai_act_status = "✓ COMPLIANT" if data.compliance_status.ai_act_compliant else "✗ NON-COMPLIANT"

            elements.append(Paragraph(f"<b>GDPR:</b> {gdpr_status}", styles["Normal"]))
            elements.append(Paragraph(f"<b>EU AI Act:</b> {ai_act_status}", styles["Normal"]))

            if data.compliance_status.fairness_score:
                elements.append(
                    Paragraph(
                        f"<b>Fairness Score:</b> {data.compliance_status.fairness_score * 100:.1f}%",
                        styles["Normal"],
                    )
                )

        return elements

    def _build_risk_factors(self, styles, data):
        """Build risk factors section."""
        elements = []
        elements.append(Paragraph("Identified Risk Factors", styles["CustomHeading"]))

        for i, factor in enumerate(data.risk_factors, 1):
            elements.append(Paragraph(f"<b>{i}.</b> {factor}", styles["Normal"]))

        return elements

    def _build_recommendations(self, styles, data):
        """Build recommendations section."""
        elements = []
        elements.append(Paragraph("Recommendations", styles["CustomHeading"]))

        for i, rec in enumerate(data.recommendations, 1):
            elements.append(Paragraph(f"<b>{i}.</b> {rec}", styles["Normal"]))

        return elements

    def _build_footer(self, styles, data):
        """Build report footer."""
        elements = []
        elements.append(
            Paragraph(
                "This report was generated by FairMind AI Governance Platform.",
                styles["Normal"],
            )
        )
        elements.append(
            Paragraph(f"Report ID: {data.timestamp.replace(':', '').replace('-', '')}",
                      styles["Normal"])
        )
        return elements


# ─── Python-DOCX DOCX Generator ─────────────────────────────────────────────

class DOCXReportGenerator:
    """Generate professional DOCX reports using python-docx."""

    def generate(self, data: BiasReportData) -> bytes:
        """Generate DOCX report and return as bytes."""
        doc = Document()

        # Set up document margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(0.75)
            section.bottom_margin = Inches(0.75)
            section.left_margin = Inches(0.75)
            section.right_margin = Inches(0.75)

        # Header
        self._add_header(doc, data)

        # Executive Summary
        self._add_executive_summary(doc, data)

        # Evaluation Summary
        self._add_evaluation_summary(doc, data)

        # Compliance Status
        self._add_compliance_status(doc, data)

        # Risk Factors
        if data.risk_factors:
            self._add_risk_factors(doc, data)

        # Recommendations
        if data.recommendations:
            self._add_recommendations(doc, data)

        # Footer
        self._add_footer(doc, data)

        # Save to bytes
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()

    def _add_header(self, doc, data):
        """Add header with FairMind branding."""
        title = doc.add_paragraph()
        title.style = "Heading 1"
        run = title.add_run("FairMind")
        run.font.size = Pt(28)
        run.font.bold = True
        run.font.color.rgb = RGBColor(31, 41, 55)

        subtitle = doc.add_paragraph("AI Governance & Bias Detection Platform")
        subtitle_run = subtitle.runs[0]
        subtitle_run.font.size = Pt(10)
        subtitle_run.font.color.rgb = RGBColor(107, 114, 128)

        report_title = doc.add_paragraph("Bias Evaluation Report")
        report_title.style = "Heading 2"

        gen_date = doc.add_paragraph(
            f"Generated: {datetime.fromisoformat(data.timestamp).strftime('%B %d, %Y at %H:%M UTC')}"
        )
        gen_date_run = gen_date.runs[0]
        gen_date_run.font.size = Pt(10)
        gen_date_run.font.color.rgb = RGBColor(107, 114, 128)

        doc.add_paragraph()  # Spacer

    def _add_executive_summary(self, doc, data):
        """Add executive summary."""
        heading = doc.add_paragraph("Executive Summary")
        heading.style = "Heading 2"

        risk_p = doc.add_paragraph()
        risk_p.add_run("Overall Risk Level: ").bold = True
        risk_run = risk_p.add_run(data.overall_risk.upper())
        risk_colors = {
            "low": RGBColor(34, 197, 94),
            "medium": RGBColor(234, 179, 8),
            "high": RGBColor(249, 115, 22),
            "critical": RGBColor(239, 68, 68),
        }
        risk_run.font.color.rgb = risk_colors.get(data.overall_risk, RGBColor(107, 114, 128))
        risk_run.font.bold = True

        if data.evaluation_summary:
            bias_rate = data.evaluation_summary.overall_bias_rate * 100
            pass_rate = (data.evaluation_summary.tests_passed / max(1, data.evaluation_summary.total_tests)) * 100
            metrics = doc.add_paragraph(
                f"Bias Rate: {bias_rate:.1f}% | Tests Passed: {data.evaluation_summary.tests_passed}/{data.evaluation_summary.total_tests} ({pass_rate:.0f}%)"
            )
            metrics_run = metrics.runs[0]
            metrics_run.font.size = Pt(10)

        doc.add_paragraph()  # Spacer

    def _add_evaluation_summary(self, doc, data):
        """Add evaluation summary table."""
        heading = doc.add_paragraph("Evaluation Summary")
        heading.style = "Heading 2"

        if data.evaluation_summary:
            table = doc.add_table(rows=5, cols=2)
            table.style = "Light Grid Accent 1"

            # Header row
            table.cell(0, 0).text = "Metric"
            table.cell(0, 1).text = "Value"

            # Data rows
            table.cell(1, 0).text = "Total Tests"
            table.cell(1, 1).text = str(data.evaluation_summary.total_tests)

            table.cell(2, 0).text = "Tests Passed"
            table.cell(2, 1).text = str(data.evaluation_summary.tests_passed)

            table.cell(3, 0).text = "Tests Failed"
            table.cell(3, 1).text = str(data.evaluation_summary.tests_failed)

            table.cell(4, 0).text = "Overall Bias Rate"
            table.cell(4, 1).text = f"{data.evaluation_summary.overall_bias_rate * 100:.1f}%"

        doc.add_paragraph()  # Spacer

    def _add_compliance_status(self, doc, data):
        """Add compliance status."""
        heading = doc.add_paragraph("Compliance Status")
        heading.style = "Heading 2"

        if data.compliance_status:
            gdpr_p = doc.add_paragraph()
            gdpr_p.add_run("GDPR: ").bold = True
            gdpr_status = "✓ COMPLIANT" if data.compliance_status.gdpr_compliant else "✗ NON-COMPLIANT"
            gdpr_p.add_run(gdpr_status)

            ai_act_p = doc.add_paragraph()
            ai_act_p.add_run("EU AI Act: ").bold = True
            ai_act_status = "✓ COMPLIANT" if data.compliance_status.ai_act_compliant else "✗ NON-COMPLIANT"
            ai_act_p.add_run(ai_act_status)

            if data.compliance_status.fairness_score:
                fairness_p = doc.add_paragraph()
                fairness_p.add_run("Fairness Score: ").bold = True
                fairness_p.add_run(f"{data.compliance_status.fairness_score * 100:.1f}%")

        doc.add_paragraph()  # Spacer

    def _add_risk_factors(self, doc, data):
        """Add risk factors."""
        heading = doc.add_paragraph("Identified Risk Factors")
        heading.style = "Heading 2"

        for i, factor in enumerate(data.risk_factors, 1):
            p = doc.add_paragraph(factor, style="List Number")

        doc.add_paragraph()  # Spacer

    def _add_recommendations(self, doc, data):
        """Add recommendations."""
        heading = doc.add_paragraph("Recommendations")
        heading.style = "Heading 2"

        for i, rec in enumerate(data.recommendations, 1):
            p = doc.add_paragraph(rec, style="List Number")

        doc.add_paragraph()  # Spacer

    def _add_footer(self, doc, data):
        """Add footer."""
        footer_para = doc.add_paragraph("This report was generated by FairMind AI Governance Platform.")
        footer_run = footer_para.runs[0]
        footer_run.font.size = Pt(9)
        footer_run.font.color.rgb = RGBColor(107, 114, 128)

        report_id = doc.add_paragraph(f"Report ID: {data.timestamp.replace(':', '').replace('-', '')}")
        report_id_run = report_id.runs[0]
        report_id_run.font.size = Pt(9)
        report_id_run.font.color.rgb = RGBColor(107, 114, 128)


# ─── Report Service (Facade) ────────────────────────────────────────────────

class ReportGenerator:
    """Facade for report generation (PDF and DOCX)."""

    def __init__(self):
        self.pdf_gen = PDFReportGenerator()
        self.docx_gen = DOCXReportGenerator()

    def generate_pdf(self, data: BiasReportData) -> bytes:
        """Generate PDF report."""
        return self.pdf_gen.generate(data)

    def generate_docx(self, data: BiasReportData) -> bytes:
        """Generate DOCX report."""
        return self.docx_gen.generate(data)
